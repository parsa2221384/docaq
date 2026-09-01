from io import BytesIO
from unittest.mock import patch

from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import TestCase
from docx import Document as DocxDocument

from rest_framework import status
from rest_framework.test import APITestCase

from .models import Document
from .utils import extract_text_from_docx


def create_test_docx():
    document = DocxDocument()

    document.add_paragraph(
        "Django is a Python web framework."
    )
    document.add_paragraph(
        "Django provides an ORM."
    )

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    return SimpleUploadedFile(
        "test.docx",
        buffer.read(),
        content_type=(
            "application/vnd.openxmlformats-officedocument"
            ".wordprocessingml.document"
        ),
    )


class DocumentModelTests(TestCase):
    def test_create_document(self):
        document = Document.objects.create(
            title="Test Document",
            content="Test content.",
        )

        self.assertEqual(
            document.title,
            "Test Document",
        )

        self.assertEqual(
            document.content,
            "Test content.",
        )

        self.assertIsNotNone(
            document.created_at,
        )

    def test_string_representation(self):
        document = Document.objects.create(
            title="Django",
            content="Test content.",
        )

        self.assertEqual(
            str(document),
            "Django",
        )


class DocumentExtractionTests(TestCase):
    def test_extract_text_from_docx(self):
        uploaded_file = create_test_docx()

        text = extract_text_from_docx(
            uploaded_file
        )

        self.assertIn(
            "Django is a Python web framework.",
            text,
        )

        self.assertIn(
            "Django provides an ORM.",
            text,
        )


class DocumentAPITests(APITestCase):
    @patch("documents.signals.index_document")
    def test_create_document(
        self,
        mock_index_document,
    ):
        response = self.client.post(
            "/api/documents/",
            {
                "title": "Test DOCX",
                "file": create_test_docx(),
            },
            format="multipart",
        )

        self.assertEqual(
            response.status_code,
            status.HTTP_201_CREATED,
        )

        self.assertEqual(
            Document.objects.count(),
            1,
        )

        document = Document.objects.first()

        self.assertEqual(
            document.title,
            "Test DOCX",
        )

        self.assertTrue(
            document.content,
        )

        mock_index_document.assert_called_once()

    def test_list_documents(self):
        Document.objects.create(
            title="Django",
            content="Django content.",
        )

        response = self.client.get(
            "/api/documents/"
        )

        self.assertEqual(
            response.status_code,
            status.HTTP_200_OK,
        )

        self.assertEqual(
            len(response.data),
            1,
        )

    @patch("documents.signals.index_document")
    def test_invalid_file_rejected(
        self,
        mock_index_document,
    ):
        uploaded_file = SimpleUploadedFile(
            "test.pdf",
            b"fake pdf content",
            content_type="application/pdf",
        )

        response = self.client.post(
            "/api/documents/",
            {
                "title": "Invalid",
                "file": uploaded_file,
            },
            format="multipart",
        )

        self.assertEqual(
            response.status_code,
            status.HTTP_400_BAD_REQUEST,
        )

        self.assertIn(
            "file",
            response.data,
        )

        mock_index_document.assert_not_called()

    @patch("documents.signals.vector_store")
    def test_delete_document(
        self,
        mock_vector_store,
    ):
        document = Document.objects.create(
            title="Django",
            content="Django content.",
        )

        document_id = document.id

        document.delete()

        self.assertFalse(
            Document.objects.filter(
                id=document_id
            ).exists()
        )

        mock_vector_store.delete.assert_called_once_with(
            where={
                "document_id": document_id,
            }
        )


class DocumentUpdateTests(APITestCase):
    @patch("documents.signals.index_document")
    def test_update_document(
        self,
        mock_index_document,
    ):
        document = Document.objects.create(
            title="Old Title",
            content="Old content.",
        )
        mock_index_document.reset_mock()
        response = self.client.patch(
            f"/api/documents/{document.id}/",
            {
                "title": "New Title",
            },
            format="multipart",
        )

        self.assertEqual(
            response.status_code,
            status.HTTP_200_OK,
        )

        document.refresh_from_db()

        self.assertEqual(
            document.title,
            "New Title",
        )

        self.assertEqual(
            document.content,
            "Old content.",
        )

        mock_index_document.assert_called_once()


class DocumentSignalTests(TestCase):
    @patch("documents.signals.vector_store")
    def test_delete_document_removes_vectors(
        self,
        mock_vector_store,
    ):
        document = Document.objects.create(
            title="Django",
            content="Django content.",
        )

        document_id = document.id

        document.delete()

        mock_vector_store.delete.assert_called_once_with(
            where={
                "document_id": document_id,
            }
        )

