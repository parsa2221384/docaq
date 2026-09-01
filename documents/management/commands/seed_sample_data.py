from io import BytesIO

from django.core.files.base import ContentFile
from django.core.management.base import BaseCommand
from docx import Document as DocxDocument

from documents.models import Document
from documents.rag.vector_store import index_document


SAMPLES = [
    {
        "title": "Django Basics",
        "paragraphs": [
            "Django is a Python web framework.",
            "Django provides an ORM for interacting with databases.",
            "Django includes an administration interface.",
        ],
    },
    {
        "title": "RAG Basics",
        "paragraphs": [
            "Retrieval-Augmented Generation combines retrieval "
            "with language model generation.",
            "A retriever finds relevant chunks from a knowledge base.",
            "The retrieved context is provided to the language model.",
        ],
    },
    {
        "title": "Java OOP",
        "paragraphs": [
            "Java supports inheritance between classes.",
            "An overridden method in a child class can replace "
            "the implementation of a parent class.",
            "Static methods belong to the class rather than an instance.",
        ],
    },
]


class Command(BaseCommand):
    help = "Create sample DOCX documents and index them."

    def handle(self, *args, **options):
        for sample in SAMPLES:
            if Document.objects.filter(
                title=sample["title"]
            ).exists():
                self.stdout.write(
                    f"Skipping existing: {sample['title']}"
                )
                continue

            docx = DocxDocument()

            for paragraph in sample["paragraphs"]:
                docx.add_paragraph(paragraph)

            buffer = BytesIO()
            docx.save(buffer)
            buffer.seek(0)

            document = Document(
                title=sample["title"],
                content="",
            )

            document.file.save(
                f"{sample['title'].lower().replace(' ', '_')}.docx",
                ContentFile(buffer.read()),
                save=False,
            )

            document.content = "\n".join(
                sample["paragraphs"]
            )

            document.save()

            index_document(document)

            self.stdout.write(
                self.style.SUCCESS(
                    f"Created: {document.title}"
                )
            )

        self.stdout.write(
            self.style.SUCCESS(
                "Sample data creation completed."
            )
        )